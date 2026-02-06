import pypdf
import docx
import io
import pytesseract
import json
import re
import ast
import string
from PIL import Image

try:
    from pdf2image import convert_from_bytes
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

def clean_extracted_text(text):
    """
    Removes non-printable characters and normalizes whitespace.
    Essential for ensuring LLM prompts remain clean and JSON isn't broken.
    """
    if not text:
        return ""
    # Filter out non-printable characters (keep standard ASCII and basic punctuation)
    printable = set(string.printable)
    text = "".join(filter(lambda x: x in printable or x.isspace(), text))

    # Normalize line breaks: replace multiple newlines with double newlines
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()

def calculate_text_quality(text):
    """
    Returns a heuristic score (0-100) based on text legibility.
    Low scores usually indicate bad PDF encoding or OCR failure.
    """
    if not text or len(text.strip()) < 50:
        return 0

    words = text.split()
    avg_word_len = sum(len(w) for w in words) / len(words) if words else 0

    # Heuristics for "Garbage" detection
    score = 100

    # 1. Extremely long average words usually mean a space-encoding error in PDF
    if avg_word_len > 18: score -= 40
    if avg_word_len > 30: score -= 60

    # 2. Too few letters (high symbol/number density)
    alpha_count = sum(1 for c in text if c.isalpha())
    ratio = alpha_count / len(text) if len(text) > 0 else 0
    if ratio < 0.5: score -= 30
    if ratio < 0.3: score -= 50

    return max(0, score)

def extract_text_from_docx(file_bytes):
    """
    Extracts text from DOCX files, processing paragraphs and tables.
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        return clean_extracted_text("\n".join(full_text))
    except Exception as e:
        return f"Error reading DOCX: {e}"

def extract_text_from_pdf(file_bytes, use_ocr=False, log_callback=None):
    """
    Advanced PDF extraction with native-to-OCR fallback.
    Triggers OCR if:
    1. Native extraction crashes completely.
    2. Any single page fails (partial failure).
    3. Extracted text quality is low.
    """
    def log(msg):
        if log_callback:
            log_callback(msg)
        print(f"[DEBUG] PDF: {msg}")

    text = ""
    native_failed = False
    partial_failure = False # Track if specific pages failed

    # 1. Try Native Extraction first
    try:
        log("Attempting native PDF extraction...")
        pdf_reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        for i, page in enumerate(pdf_reader.pages):
            try:
                content = page.extract_text()
                if content:
                    text += content + "\n"
            except Exception as e:
                log(f"Page {i+1} extraction warning: {e}")
                partial_failure = True # Mark that at least one page failed
                continue
    except Exception as e:
        log(f"Native extraction crashed: {e}")
        native_failed = True

    # 2. Assess Quality
    quality = calculate_text_quality(text)
    log(f"Native Quality Score: {quality}/100")

    # 3. OCR Fallback Logic
    # Trigger if:
    # a) Native extraction crashed completely
    # b) Any specific page failed (partial_failure) <-- NEW CHECK
    # c) Text is too short (< 150 chars)
    # d) Quality is garbage (< 50)
    should_use_ocr = (native_failed or partial_failure or len(text.strip()) < 150 or quality < 50) and use_ocr

    if should_use_ocr:
        reason = "Partial Page Failure" if partial_failure else "Low Quality/Crash"

        if not PDF2IMAGE_AVAILABLE:
            log(f"OCR needed ({reason}) but dependencies missing.")
            return text if len(text.strip()) > 50 else "[Error: OCR required but dependencies missing]"

        log(f"⚠️ {reason} detected. Triggering OCR Fallback...")
        try:
            images = convert_from_bytes(file_bytes, dpi=300)
            ocr_text = ""
            for i, img in enumerate(images):
                log(f"OCR Page {i+1}...")
                page_text = pytesseract.image_to_string(img, config='--psm 3')
                ocr_text += page_text + "\n"

            final_text = clean_extracted_text(ocr_text)
            log(f"OCR Complete. Final Score: {calculate_text_quality(final_text)}/100")
            return final_text
        except Exception as e:
            log(f"OCR Failed: {e}")
            return text

    return clean_extracted_text(text)

def clean_json_response(text):
    """
    Robust extraction of JSON from LLM markdown response.
    Uses hex escape sequences for backticks (\x60) to prevent breaking the editor UI.
    """
    if not text or not isinstance(text, str):
        return None

    try:
        # Define backticks using hex escape to avoid closing the parent markdown block
        backticks = "\x60\x60\x60"

        # 1. Look for fenced JSON code block using hex-escaped pattern
        pattern = rf"{backticks}(?:json)?\s*(.*?)\s*{backticks}"
        match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)

        if match:
            text_content = match.group(1).strip()
        else:
            # 2. Fallback: Find the outermost curly braces or square brackets
            start = text.find('{')
            end = text.rfind('}')
            if start != -1 and end != -1:
                text_content = text[start:end+1].strip()
            else:
                # Try finding list brackets if object failed
                start_list = text.find('[')
                end_list = text.rfind(']')
                if start_list != -1 and end_list != -1 and end_list > start_list:
                    text_content = text[start_list:end_list+1].strip()
                else:
                    return None

        # Normalize smart quotes and non-standard characters
        text_content = text_content.replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"')

        # 3. Standard JSON decode
        try:
            return json.loads(text_content)
        except json.JSONDecodeError:
            # 4. Final attempt: literal_eval for Python-style dictionaries/lists
            try:
                return ast.literal_eval(text_content)
            except Exception:
                return None

    except Exception:
        return None

def _is_list_of_strings(value):
    return isinstance(value, list) and all(isinstance(v, str) for v in value)

def _is_list_of_dicts(value):
    return isinstance(value, list) and all(isinstance(v, dict) for v in value)

def validate_jd_schema(data):
    errors = []
    if not isinstance(data, dict):
        return False, ["JD data is not an object"]

    required = {
        "role_title": str,
        "must_have_skills": list,
        "nice_to_have_skills": list,
        "min_years_experience": int,
        "education_requirements": list,
        "domain_knowledge": list,
        "soft_skills": list,
        "key_responsibilities": list,
    }

    for key, typ in required.items():
        if key not in data:
            errors.append(f"Missing key: {key}")
            continue
        if not isinstance(data[key], typ):
            errors.append(f"Invalid type for {key}")

    # List element sanity checks
    for key in [
        "must_have_skills",
        "nice_to_have_skills",
        "education_requirements",
        "domain_knowledge",
        "soft_skills",
        "key_responsibilities",
    ]:
        if key in data and not _is_list_of_strings(data[key]):
            errors.append(f"{key} must be a list of strings")

    return len(errors) == 0, errors

def validate_resume_profile_schema(data):
    errors = []
    if not isinstance(data, dict):
        return False, ["Resume profile is not an object"]

    required = {
        "candidate_name": str,
        "email": str,
        "phone": str,
        "extracted_skills": list,
        "years_experience": int,
        "education_summary": str,
        "domain_experience": list,
        "work_history": list,
    }

    for key, typ in required.items():
        if key not in data:
            errors.append(f"Missing key: {key}")
            continue
        if not isinstance(data[key], typ):
            errors.append(f"Invalid type for {key}")

    if "extracted_skills" in data and not _is_list_of_strings(data["extracted_skills"]):
        errors.append("extracted_skills must be a list of strings")
    if "domain_experience" in data and not _is_list_of_strings(data["domain_experience"]):
        errors.append("domain_experience must be a list of strings")
    if "work_history" in data and not _is_list_of_dicts(data["work_history"]):
        errors.append("work_history must be a list of objects")

    return len(errors) == 0, errors

def validate_criterion_schema(data):
    errors = []
    if not isinstance(data, dict):
        return False, ["Criterion is not an object"]
    for key in ["requirement", "status", "evidence"]:
        if key not in data:
            errors.append(f"Missing key: {key}")
    if "status" in data and data.get("status") not in ["Met", "Partial", "Missing"]:
        errors.append("Invalid status value")
    return len(errors) == 0, errors

def validate_bulk_criteria_schema(data):
    if not isinstance(data, list):
        return False, ["Bulk criteria is not a list"]
    errors = []
    for idx, item in enumerate(data):
        ok, err = validate_criterion_schema(item)
        if not ok:
            errors.append(f"Item {idx} invalid: {', '.join(err)}")
    return len(errors) == 0, errors

def validate_standard_output_schema(data):
    errors = []
    if not isinstance(data, dict):
        return False, ["Standard output is not an object"]
    for key in ["candidate_name", "match_score", "decision", "reasoning", "missing_skills", "match_details"]:
        if key not in data:
            errors.append(f"Missing key: {key}")
    if "match_score" in data and not isinstance(data["match_score"], int):
        errors.append("match_score must be int")
    if "missing_skills" in data and not _is_list_of_strings(data["missing_skills"]):
        errors.append("missing_skills must be a list of strings")
    if "match_details" in data and not _is_list_of_dicts(data["match_details"]):
        errors.append("match_details must be a list of objects")
    return len(errors) == 0, errors
