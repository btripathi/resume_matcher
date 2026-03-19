"""Tests for document_utils.py."""

import io

from document_utils import clean_extracted_text, clean_json_response, extract_text_from_docx


def test_clean_extracted_text_strips_nonprintable():
    text = "Hello\x00World\x01\nGood"
    result = clean_extracted_text(text)
    assert "\x00" not in result
    assert "\x01" not in result
    assert "Hello" in result
    assert "Good" in result


def test_clean_extracted_text_normalizes_blank_lines():
    text = "Line1\n\n\n\n\nLine2"
    result = clean_extracted_text(text)
    assert result == "Line1\n\nLine2"


def test_clean_extracted_text_empty():
    assert clean_extracted_text("") == ""
    assert clean_extracted_text(None) == ""


def test_extract_text_from_docx_minimal():
    """Create a minimal docx in memory and extract its text."""
    import docx
    doc = docx.Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("Second paragraph")
    buf = io.BytesIO()
    doc.save(buf)
    result = extract_text_from_docx(buf.getvalue())
    assert "Hello World" in result
    assert "Second paragraph" in result


def test_clean_json_response_array():
    raw = '[{"status": "Met"}, {"status": "Missing"}]'
    result = clean_json_response(raw)
    assert isinstance(result, list)
    assert len(result) == 2
    assert result[0]["status"] == "Met"


def test_clean_json_response_with_surrounding_text():
    raw = 'Here is the result:\n{"score": 42}\nEnd of output.'
    result = clean_json_response(raw)
    assert result == {"score": 42}
